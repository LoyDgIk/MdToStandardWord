# -*- coding: utf-8 -*-
"""把 model.StandardDoc 构建进团体标准模板，输出标准文本 Word。

核心思路：复制模板 .docx，逐节回填。
- 封面字段：按样式名定位占位段落，替换 run 文本（保留格式）。
- 前言/引言/正文/参考文献/索引：每节由一个"承载 sectPr 的段落"作为终止段；
  清空该节内容段落（保留终止段以维持分节符与页眉页脚），再在终止段前插入新内容。
- 章条编号全部依赖模板的样式联动多级列表自动生成，标题只写文本、不写编号。
"""

from __future__ import annotations

import hashlib
import io
import copy
import os
import re
import shutil
import zipfile
from typing import List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from . import boilerplate as bp
from . import model
from . import resources
from . import styles as S


# --------------------------------------------------------------------------- #
# 低层 OXML 帮助
# --------------------------------------------------------------------------- #
def _carries_sectpr(p: Paragraph) -> bool:
    ppr = p._p.find(qn("w:pPr"))
    if ppr is None:
        return False
    return ppr.find(qn("w:sectPr")) is not None


def _set_numbering(para: Paragraph, num_id: int, ilvl: int):
    """在段落上显式设置列表编号（numId/ilvl），覆盖样式默认。

    num_id=0 表示关闭编号。用于让"无标题条"接入正文章条同一多级列表(numId=2)，
    从而与章/条编号连续（经 Word 验证：标准文件_X无标题 + numId=2 即可同步）。
    """
    p = para._p
    ppr = p.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        p.insert(0, ppr)
    old = ppr.find(qn("w:numPr"))
    if old is not None:
        ppr.remove(old)
    numpr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    numid_el = OxmlElement("w:numId")
    numid_el.set(qn("w:val"), str(num_id))
    numpr.append(ilvl_el)
    numpr.append(numid_el)
    pstyle = ppr.find(qn("w:pStyle"))
    if pstyle is not None:
        pstyle.addnext(numpr)
    else:
        ppr.insert(0, numpr)


def _style_numbering(doc, style_name: str):
    """读取模板样式自带的 numPr，返回 (numId, ilvl)。"""
    try:
        style = doc.styles[style_name]
    except KeyError:
        return None
    ppr = style.element.find(qn("w:pPr"))
    if ppr is None:
        return None
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        return None
    numid = numpr.find(qn("w:numId"))
    if numid is None or numid.get(qn("w:val")) is None:
        return None
    ilvl = numpr.find(qn("w:ilvl"))
    return int(numid.get(qn("w:val"))), int(ilvl.get(qn("w:val")) if ilvl is not None and ilvl.get(qn("w:val")) is not None else 0)


def _abstract_num_id_for_num(doc, num_id: int) -> Optional[int]:
    """返回 numId 对应的 abstractNumId。"""
    numbering = doc.part.numbering_part.element
    for num in numbering.findall(qn("w:num")):
        if num.get(qn("w:numId")) != str(num_id):
            continue
        abstract = num.find(qn("w:abstractNumId"))
        if abstract is None or abstract.get(qn("w:val")) is None:
            return None
        return int(abstract.get(qn("w:val")))
    return None


def _next_numbering_id(doc) -> int:
    numbering = doc.part.numbering_part.element
    ids = []
    for num in numbering.findall(qn("w:num")):
        val = num.get(qn("w:numId"))
        if val is not None and val.isdigit():
            ids.append(int(val))
    return (max(ids) if ids else 0) + 1


def _new_numbering_instance_from_style(doc, style_name: str) -> Optional[int]:
    """基于模板样式的 abstractNum 新建一个 numId，使独立列表重新编号。"""
    style_numbering = _style_numbering(doc, style_name)
    if style_numbering is None:
        return None
    base_num_id, _ = style_numbering
    abstract_num_id = _abstract_num_id_for_num(doc, base_num_id)
    if abstract_num_id is None:
        return None

    new_num_id = _next_numbering_id(doc)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), str(abstract_num_id))
    num.append(abstract)
    for ilvl in range(9):
        lvl_override = OxmlElement("w:lvlOverride")
        lvl_override.set(qn("w:ilvl"), str(ilvl))
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), "1")
        lvl_override.append(start_override)
        num.append(lvl_override)
    doc.part.numbering_part.element.append(num)
    return new_num_id


def _set_numbering_from_style(para: Paragraph, doc, style_name: str,
                              num_id_override: Optional[int] = None):
    """显式套用模板样式的编号定义，避免 Word 忽略样式级列表缩进。"""
    numbering = _style_numbering(doc, style_name)
    if numbering is None:
        return
    num_id, ilvl = numbering
    if num_id_override is not None:
        num_id = num_id_override
    _set_numbering(para, num_id, ilvl)


def _first_existing_path(candidates: List[str], label: str) -> str:
    for path in candidates:
        if path and os.path.isfile(path):
            return path
    raise FileNotFoundError("找不到%s：%s" % (label, "；".join(candidates)))


def _resolve_kind(kind: str, meta: model.Meta) -> str:
    if kind not in ("auto", "group", "national"):
        raise ValueError("kind 只能是 auto、group 或 national。")
    if kind != "auto":
        return kind
    standard_type = (meta.standard_type or "").strip()
    number = (meta.number or "").strip().upper()
    if "国家" in standard_type or number.startswith("GB"):
        return "national"
    return "group"


def _default_cover_path(kind: str) -> str:
    if kind == "national":
        return _first_existing_path(
            resources.template_candidates("cover_national.docx", "template_national.docx"),
            "国家标准封面蓝图",
        )
    return _first_existing_path(
        resources.template_candidates("cover_group.docx", "template_group.docx"),
        "团体标准封面蓝图",
    )


def _cover_blueprint_from_source(source_path: str, output_path: str):
    """从完整模板截取封面页为蓝图，保留原封面的图片、关系和首个分节符。"""
    from lxml import etree

    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ns = {"w": w_ns}
    with zipfile.ZipFile(source_path, "r") as zin, zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/document.xml":
                root = etree.fromstring(data)
                body = root.find("w:body", namespaces=ns)
                children = list(body)
                final_sect = None
                for child in children:
                    if child.tag == "{%s}sectPr" % w_ns:
                        final_sect = child
                        break

                keep = []
                for child in children:
                    if child.tag == "{%s}sectPr" % w_ns:
                        continue
                    keep.append(child)
                    if child.tag == "{%s}p" % w_ns and child.find("w:pPr/w:sectPr", namespaces=ns) is not None:
                        break

                for child in list(body):
                    body.remove(child)
                for child in keep:
                    body.append(child)
                if final_sect is not None:
                    body.append(final_sect)
                data = etree.tostring(
                    root,
                    xml_declaration=True,
                    encoding="UTF-8",
                    standalone=True,
                )
            zout.writestr(item, data)


def _copy_cover_base(cover_path: str, output_path: str):
    """复制封面蓝图；若传入完整模板，则运行时截取封面，不带入正文占位章。"""
    # 正式蓝图文件名固定为 cover_*.docx。其它 docx 来源按完整模板处理并截取首节。
    name = os.path.basename(cover_path).lower()
    if name.startswith("cover_"):
        shutil.copyfile(cover_path, output_path)
    else:
        _cover_blueprint_from_source(cover_path, output_path)


def _read_cover_end_line_image(docx_path: str, kind: str) -> bytes:
    """读取封面蓝图包内自带的标准结束线图片。"""
    if kind == "national":
        candidates = [
            "word/media/image3.jpg",
            "word/media/image3.jpeg",
        ]
    else:
        candidates = [
            "word/media/image1.jpeg",
            "word/media/image1.jpg",
        ]
    with zipfile.ZipFile(docx_path, "r") as zf:
        for name in candidates:
            try:
                return zf.read(name)
            except KeyError:
                continue
    raise FileNotFoundError("封面蓝图缺少标准结束线图片：%s" % "；".join(candidates))


def _reset_counters():
    _COUNTER.table = 0
    _COUNTER.figure = 0
    _COUNTER.bm = 1000
    _COUNTER.seq_scope_counts = {}


def _configure_standard_styles(doc):
    # 模板自带的破折号列项/参考文献缩进偏大。这里仍使用模板样式和编号，
    # 只把编号级别缩进收敛到“自然段首行两字”附近。
    _clear_style_indent(doc, S.S_LIST_DASH)
    _clear_style_indent(doc, "标准文件_破折号列项（二级）")
    _fix_numbering_style_indent(doc, S.S_LIST_DASH, left_twips=600, hanging_twips=200)
    _fix_numbering_style_indent(doc, "标准文件_破折号列项（二级）", left_twips=920, hanging_twips=200)
    _fix_numbering_style_indent(doc, S.S_REF_ITEM, left_twips=620, hanging_twips=420)



# --------------------------------------------------------------------------- #
# 域 / 书签（SEQ 自动编号、REF 交叉引用）
# ---------------------------------------------------------------------------
# SEQ 前缀与可见题注标签保持一致，避免 Word 只把数字识别为题注范围。
# 表/图/公式都使用可见 SEQ；模板标题样式只负责外观，不再负责编号。
# 书签：
#   _Ref... → Word 原生隐藏书签，分别围住编号、标签+编号、整项题注等范围。
# --------------------------------------------------------------------------- #

SEQ_TABLE    = "表"
SEQ_FIGURE   = "图"
SEQ_EQUATION = "公式"

def _bm_name(anchor_id: str) -> str:
    """把任意 id（可含中文）映射为合法的 Word 书签名。

    纯 ASCII 标识（如标准号 "GB 5749"）生成可读名 "Ref_GB5749"，便于在
    Word 原生交叉引用对话框里识别；含中文等则用哈希，保证有效且稳定。
    """
    norm = anchor_id.replace(" ", "").strip()
    ascii_safe = re.sub(r"[^0-9A-Za-z]+", "_", norm)
    if norm and re.match(r"^[0-9A-Za-z/._\-]+$", norm) and ascii_safe:
        return ("Ref_" + ascii_safe)[:40]
    return "Ref_" + hashlib.md5(norm.encode("utf-8")).hexdigest()[:16]


def _native_ref_name(ref_type: str, local_id: str, mode: str) -> str:
    """生成 Word 原生交叉引用常用的隐藏书签名 `_Ref#########`。"""
    key = "%s:%s:%s" % (ref_type, local_id, mode)
    n = int(hashlib.md5(key.encode("utf-8")).hexdigest()[:8], 16) % 900000000
    return "_Ref" + str(100000000 + n)


def _bookmark_start(para: Paragraph, name: str, bid: int):
    bs = OxmlElement("w:bookmarkStart")
    bs.set(qn("w:id"), str(bid))
    bs.set(qn("w:name"), name)
    para._p.append(bs)


def _bookmark_end(para: Paragraph, bid: int):
    be = OxmlElement("w:bookmarkEnd")
    be.set(qn("w:id"), str(bid))
    para._p.append(be)


def _add_ref_bookmark(para: Paragraph, bookmark_name: str, display_text: str = "?",
                      charformat: bool = False):
    """插入 REF 域，引用指定书签。display_text 是域更新前的占位结果。"""
    r = para.add_run()
    fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin"); r._r.append(fb)
    r = para.add_run()
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
    fmt = " \\* CHARFORMAT" if charformat else ""
    it.text = " REF %s \\h%s " % (bookmark_name, fmt)
    r._r.append(it)
    r = para.add_run()
    fs = OxmlElement("w:fldChar"); fs.set(qn("w:fldCharType"), "separate"); r._r.append(fs)
    para.add_run(display_text)
    r = para.add_run()
    fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end"); r._r.append(fe)


def _hyperlink_run_text(text: str):
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    if text[:1].isspace() or text[-1:].isspace():
        t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _hyperlink_run_fld_char(kind: str):
    r = OxmlElement("w:r")
    fc = OxmlElement("w:fldChar")
    fc.set(qn("w:fldCharType"), kind)
    r.append(fc)
    return r


def _hyperlink_run_instr(text: str):
    r = OxmlElement("w:r")
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = text
    r.append(it)
    return r


def _add_hyperlinked_ref_bookmark(para: Paragraph, hyperlink_anchor: str,
                                  ref_bookmark: str, prefix: str = "",
                                  suffix: str = "", display_text: str = "?",
                                  charformat: bool = False):
    """插入一个外层整体可点击、内部编号可更新的 REF。"""
    link = OxmlElement("w:hyperlink")
    link.set(qn("w:anchor"), hyperlink_anchor)
    link.set(qn("w:history"), "1")
    if prefix:
        link.append(_hyperlink_run_text(prefix))
    link.append(_hyperlink_run_fld_char("begin"))
    fmt = " \\* CHARFORMAT" if charformat else ""
    link.append(_hyperlink_run_instr(" REF %s \\h%s " % (ref_bookmark, fmt)))
    link.append(_hyperlink_run_fld_char("separate"))
    link.append(_hyperlink_run_text(display_text))
    link.append(_hyperlink_run_fld_char("end"))
    if suffix:
        link.append(_hyperlink_run_text(suffix))
    para._p.append(link)


def _add_ref(para: Paragraph, anchor_id: str, suffix: str = ""):
    """插入 REF 域，引用 Ref_{anchor_id}{suffix} 书签。"""
    _add_ref_bookmark(para, _bm_name(anchor_id) + suffix)


def _add_typed_ref(para: Paragraph, ref: model.RefSpan):
    if ref.ref_type == "std":
        _add_ref(para, ref.target)
        return
    if ref.ref_type == "eq" and ref.mode in ("label", "full"):
        _add_hyperlinked_ref_bookmark(
            para,
            _native_ref_name(ref.ref_type, ref.target, "label"),
            _native_ref_name(ref.ref_type, ref.target, "num"),
            prefix="式（",
            suffix="）",
            display_text="?",
            charformat=True,
        )
        return
    _add_ref_bookmark(para, _native_ref_name(ref.ref_type, ref.target, ref.mode))


def _make_run_hidden(run_elem):
    """给一个 <w:r> 加上 <w:vanish/> 使其成为隐藏文字。"""
    rpr = run_elem.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        run_elem.insert(0, rpr)
    rpr.append(OxmlElement("w:vanish"))


def _make_paragraph_hidden(para: Paragraph):
    """隐藏整个段落，保留其中字段供 Word 交叉引用对话框识别。"""
    ppr = para._p.get_or_add_pPr()
    rpr = ppr.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        ppr.append(rpr)
    if rpr.find(qn("w:vanish")) is None:
        rpr.append(OxmlElement("w:vanish"))


def _hide_run_if_needed(run, hidden: bool):
    if hidden:
        _make_run_hidden(run._r)
    return run


def _add_seq(para: Paragraph, seq_name: str, reset: bool = False, hidden: bool = False):
    """插入可见 SEQ 域。返回（SEQ 结果 run, 字段结束 run）。"""
    r = _hide_run_if_needed(para.add_run(), hidden)
    fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin"); r._r.append(fb)
    r = _hide_run_if_needed(para.add_run(), hidden)
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
    reset_part = " \\r 1" if reset else ""
    it.text = " SEQ %s \\* ARABIC%s " % (seq_name, reset_part)
    r._r.append(it)
    r = _hide_run_if_needed(para.add_run(), hidden)
    fs = OxmlElement("w:fldChar"); fs.set(qn("w:fldCharType"), "separate"); r._r.append(fs)
    result_run = _hide_run_if_needed(para.add_run("1"), hidden)
    end_run = _hide_run_if_needed(para.add_run(), hidden)
    fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end"); end_run._r.append(fe)
    return result_run, end_run


def _bookmark_start_before(run, name: str, bid: int):
    bs = OxmlElement("w:bookmarkStart")
    bs.set(qn("w:id"), str(bid))
    bs.set(qn("w:name"), name)
    run._r.addprevious(bs)


def _bookmark_end_after(run, bid: int):
    be = OxmlElement("w:bookmarkEnd")
    be.set(qn("w:id"), str(bid))
    run._r.addnext(be)


def _next_bm_id():
    bid = _COUNTER.bm
    _COUNTER.bm += 1
    return bid


def _add_hidden_seq(para: Paragraph, seq_name: str, num_prefix: str = "", separator: str = ""):
    """插入隐藏 SEQ 域（用于表/图等有样式自动编号的场景）。

    所有相关 run 带 <w:vanish/>，不可见但被 Word 域引擎识别，
    使条目出现在"表"/"图"交叉引用域中。

    num_prefix: 前缀文字（附录用，如 "A."），插在 SEQ 之前（隐藏）。
    separator:  SEQ 结果之后的分隔符（如 "　"），插在域结束之后（隐藏）。
    返回 SEQ 结果 run（隐藏）。
    """
    if num_prefix:
        r = para.add_run(num_prefix); _make_run_hidden(r._r)
    r = para.add_run(); _make_run_hidden(r._r)
    fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin"); r._r.append(fb)
    r = para.add_run(); _make_run_hidden(r._r)
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
    it.text = " SEQ %s \\* ARABIC " % seq_name
    r._r.append(it)
    r = para.add_run(); _make_run_hidden(r._r)
    fs = OxmlElement("w:fldChar"); fs.set(qn("w:fldCharType"), "separate"); r._r.append(fs)
    result_run = para.add_run("1"); _make_run_hidden(result_run._r)
    r = para.add_run(); _make_run_hidden(r._r)
    fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end"); r._r.append(fe)
    if separator:
        r = para.add_run(separator); _make_run_hidden(r._r)
    return result_run


def _emit_caption(para, seq_name, anchor_id, title, num_prefix="", hidden=False):
    """在段落上插入 SEQ + 双书签（:a 编号 / :b 编号+标题）。

    hidden=True 时用隐藏 SEQ（表/图——样式自带列表编号），
    hidden=False 时用可见 SEQ（公式——SEQ 是唯一编号来源）。

    num_prefix: 附录前缀（如 "A."），在 SEQ 结果前显示（隐藏）。
    """
    bm_full_id = None
    bm_num_id  = None
    if anchor_id:
        bm_full_id = _COUNTER.bm; _COUNTER.bm += 1
        bm_num_id  = _COUNTER.bm; _COUNTER.bm += 1

    # 书签 full 围住：prefix + SEQ + "　" + title → {:b} 编号+标题
    if bm_full_id is not None:
        _bookmark_start(para, _bm_name(anchor_id) + "_full", bm_full_id)

    # 前缀（附录用，如 "A."）+ SEQ 域 + 分隔符
    if hidden:
        result_run = _add_hidden_seq(para, seq_name, num_prefix=num_prefix, separator="　")
    else:
        if num_prefix:
            r = para.add_run(num_prefix); _make_run_hidden(r._r)
        result_run, seq_end_run = _add_seq(para, seq_name)

    # 书签 num 仅围住 SEQ 结果（含前缀）→ {:a} 纯编号
    # 关键：bookmarkStart 插到 result_run 之前，bookmarkEnd 插到字段结束之后，使文字落入书签
    if bm_num_id is not None:
        bm_start = OxmlElement("w:bookmarkStart")
        bm_start.set(qn("w:id"), str(bm_num_id))
        bm_start.set(qn("w:name"), _bm_name(anchor_id))
        result_run._r.addprevious(bm_start)
        bm_end = OxmlElement("w:bookmarkEnd")
        bm_end.set(qn("w:id"), str(bm_num_id))
        (seq_end_run if not hidden else result_run)._r.addnext(bm_end)

    # 标题文本（visible SEQ 时加显式分隔符；hidden 已在域后插入隐藏分隔符）
    if title:
        if not hidden:
            para.add_run("　")
        para.add_run(title)

    if bm_full_id is not None:
        _bookmark_end(para, bm_full_id)


# --------------------------------------------------------------------------- #
# 缩进辅助
# --------------------------------------------------------------------------- #
# 1/100 字符 -> twips 的近似换算（按五号 10.5pt，约 210 twips/字）
_CHAR_TWIPS = 2.1
_PT_TWIPS = 20  # 1 磅 = 20 twips


def _set_ind_pt(ind, left_pt=None, hanging_pt=None, first_line_pt=None):
    """用绝对磅值重设 w:ind，清掉字符级与冲突属性。"""
    for a in ("firstLine", "firstLineChars", "hanging", "hangingChars",
              "left", "leftChars", "start", "startChars", "end", "endChars"):
        k = qn("w:" + a)
        if ind.get(k) is not None:
            del ind.attrib[k]
    if left_pt is not None:
        ind.set(qn("w:left"), str(int(round(left_pt * _PT_TWIPS))))
    if hanging_pt is not None:
        ind.set(qn("w:hanging"), str(int(round(hanging_pt * _PT_TWIPS))))
    if first_line_pt is not None:
        ind.set(qn("w:firstLine"), str(int(round(first_line_pt * _PT_TWIPS))))


def _set_ind_twips(ind, left_twips=None, hanging_twips=None, first_line_twips=None):
    """用 twips 重设 w:ind，供模板 numbering 级别微调用。"""
    for a in ("firstLine", "firstLineChars", "hanging", "hangingChars",
              "left", "leftChars", "start", "startChars", "end", "endChars"):
        k = qn("w:" + a)
        if ind.get(k) is not None:
            del ind.attrib[k]
    if left_twips is not None:
        ind.set(qn("w:left"), str(left_twips))
    if hanging_twips is not None:
        ind.set(qn("w:hanging"), str(hanging_twips))
    if first_line_twips is not None:
        ind.set(qn("w:firstLine"), str(first_line_twips))


def _set_ind(ind, left_chars=None, hanging_chars=None, first_line_chars=None):
    """重设 w:ind 的字符级缩进，清掉冲突属性。"""
    for a in ("firstLine", "firstLineChars", "hanging", "hangingChars", "left", "leftChars"):
        k = qn("w:" + a)
        if ind.get(k) is not None:
            del ind.attrib[k]
    if left_chars is not None:
        ind.set(qn("w:leftChars"), str(left_chars))
        ind.set(qn("w:left"), str(int(left_chars * _CHAR_TWIPS)))
    if hanging_chars is not None:
        ind.set(qn("w:hangingChars"), str(hanging_chars))
        ind.set(qn("w:hanging"), str(int(hanging_chars * _CHAR_TWIPS)))
    if first_line_chars is not None:
        ind.set(qn("w:firstLineChars"), str(first_line_chars))
        ind.set(qn("w:firstLine"), str(int(first_line_chars * _CHAR_TWIPS)))


def _fix_style_indent(doc, style_name, left_pt, hanging_pt):
    """改写某段落样式的缩进，用绝对磅值对齐悬挂列项。"""
    try:
        st = doc.styles[style_name]
    except KeyError:
        return
    ppr = st.element.get_or_add_pPr()
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)
    _set_ind_pt(ind, left_pt=left_pt, hanging_pt=hanging_pt)


def _clear_style_indent(doc, style_name):
    """移除段落样式自身缩进，避免与 numbering 缩进叠加。"""
    try:
        st = doc.styles[style_name]
    except KeyError:
        return
    ppr = st.element.find(qn("w:pPr"))
    if ppr is None:
        return
    ind = ppr.find(qn("w:ind"))
    if ind is not None:
        ppr.remove(ind)


def _fix_numbering_style_indent(doc, style_name, left_twips, hanging_twips):
    """按样式关联的 numbering level 修正列表/参考文献缩进。"""
    try:
        style = doc.styles[style_name]
    except KeyError:
        return
    style_id = style.element.get(qn("w:styleId"))
    if not style_id:
        return
    try:
        numbering = doc.part.numbering_part.element
    except Exception:
        return
    for lvl in numbering.findall(".//" + qn("w:lvl")):
        pstyle = lvl.find(qn("w:pStyle"))
        if pstyle is None or pstyle.get(qn("w:val")) != style_id:
            continue
        ppr = lvl.find(qn("w:pPr"))
        if ppr is None:
            ppr = OxmlElement("w:pPr")
            lvl.append(ppr)
        ind = ppr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            ppr.append(ind)
        _set_ind_twips(ind, left_twips=left_twips, hanging_twips=hanging_twips)


def _set_runs(paragraph: Paragraph, spans: List[model.Span]):
    """按 spans 给段落添加 run（加粗/斜体）；RefSpan 转换为 REF 域。"""
    from . import mathconv
    if not spans:
        return
    for sp in spans:
        if isinstance(sp, model.RefSpan):
            _add_typed_ref(paragraph, sp)
            continue
        if isinstance(sp, model.FormulaSpan):
            omath = mathconv.latex_to_omml(sp.text)
            if omath is not None:
                paragraph._p.append(omath)
            else:
                paragraph.add_run(sp.text)
            continue
        for i, piece in enumerate(sp.text.split("\n")):
            if i > 0:
                paragraph.add_run().add_break()
            if piece:
                r = paragraph.add_run(piece)
                if sp.bold:
                    r.bold = True
                if sp.italic:
                    r.italic = True
                if getattr(sp, "subscript", False):
                    r.font.subscript = True
                if getattr(sp, "superscript", False):
                    r.font.superscript = True


def _apply_style_run_properties(doc, run, style_name: str):
    """把段落样式里的 rPr 复制到 run，用于同段软换行的附录标题局部格式。"""
    try:
        style = doc.styles[style_name]
    except KeyError:
        return
    rpr = style.element.find(qn("w:rPr"))
    if rpr is None:
        return
    old = run._r.find(qn("w:rPr"))
    if old is not None:
        run._r.remove(old)
    run._r.insert(0, copy.deepcopy(rpr))


def _add_styled_runs(paragraph: Paragraph, doc, style_name: str, spans: List[model.Span]):
    from . import mathconv
    for sp in spans:
        if isinstance(sp, model.RefSpan):
            _add_typed_ref(paragraph, sp)
            continue
        if isinstance(sp, model.FormulaSpan):
            omath = mathconv.latex_to_omml(sp.text)
            if omath is not None:
                paragraph._p.append(omath)
            else:
                run = paragraph.add_run(sp.text)
                _apply_style_run_properties(doc, run, style_name)
            continue
        for i, piece in enumerate(sp.text.split("\n")):
            if i > 0:
                paragraph.add_run().add_break()
            if piece:
                run = paragraph.add_run(piece)
                _apply_style_run_properties(doc, run, style_name)
                if sp.bold:
                    run.bold = True
                if sp.italic:
                    run.italic = True
                if getattr(sp, "subscript", False):
                    run.font.subscript = True
                if getattr(sp, "superscript", False):
                    run.font.superscript = True


def _new_paragraph_before(anchor: Paragraph, doc, style_name: str,
                          text: str = "", spans: Optional[List[model.Span]] = None) -> Paragraph:
    """在 anchor 段落前插入一个套用 style_name 的新段落。"""
    new_p = OxmlElement("w:p")
    anchor._p.addprevious(new_p)
    para = Paragraph(new_p, anchor._parent)
    try:
        para.style = doc.styles[style_name]
    except KeyError:
        para.style = doc.styles[S.S_PARA]
    if spans is not None:
        _set_runs(para, spans)
    elif text:
        para.add_run(text)
    return para


def _new_numbered_style_paragraph(anchor: Paragraph, doc, style_name: str,
                                  text: str = "", spans: Optional[List[model.Span]] = None,
                                  num_id_override: Optional[int] = None) -> Paragraph:
    para = _new_paragraph_before(anchor, doc, style_name, text=text, spans=spans)
    _set_numbering_from_style(para, doc, style_name, num_id_override=num_id_override)
    return para


def _set_direct_paragraph_spacing(para: Paragraph, before: Optional[int] = None,
                                  after: Optional[int] = None):
    """设置模板标题段落保留的直接段距。"""
    if before is None and after is None:
        return
    ppr = para._p.get_or_add_pPr()
    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        ppr.append(spacing)
    if before is not None:
        spacing.set(qn("w:before"), str(before))
    if after is not None:
        spacing.set(qn("w:after"), str(after))


def _set_run_char_spacing(run, value: int, east_asia_hint: bool = False):
    rpr = run._r.get_or_add_rPr()
    if east_asia_hint:
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:hint"), "eastAsia")
    spacing = rpr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        rpr.append(spacing)
    spacing.set(qn("w:val"), str(value))


def _set_run_east_asia_hint(run):
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:hint"), "eastAsia")


def _section_title_after(title: str, kind: str) -> Optional[int]:
    front_titles = {"目次", "前言", "引言"}
    body_titles = {"参考文献", "索引"}
    if kind == "national":
        if title in front_titles:
            return 468
        if title in body_titles:
            return 156
    else:
        if title in front_titles:
            return 360
        if title in body_titles:
            return 120
    return None


def _section_title_before(title: str, kind: str) -> Optional[int]:
    # 国家标准模板的前言首页标题带直接段前距，其他标题不额外设置。
    if kind == "national" and title == "前言":
        return 900
    return None


def _section_title_char_spacing(title: str) -> tuple[Optional[int], bool]:
    if title in {"目次", "前言", "引言"}:
        return 320, False
    if title == "参考文献":
        return 105, True
    if title == "索引":
        return 210, True
    return None, False


def _emit_section_title_before(anchor: Paragraph, doc, style_name: str, title: str,
                               kind: str = "group") -> Paragraph:
    """输出与模板标题段落一致的节标题。

    模板中的"目次/前言/引言/参考文献/索引"不是单个普通 run：
    标题前 N-1 个字带字符间距，最后一个字不带字符间距；段距也有直接格式。
    """
    para = _new_paragraph_before(anchor, doc, style_name)
    _set_direct_paragraph_spacing(
        para,
        before=_section_title_before(title, kind),
        after=_section_title_after(title, kind),
    )
    char_spacing, east_asia_hint = _section_title_char_spacing(title)
    if char_spacing is None or len(title) < 2:
        para.add_run(title)
        return para

    first = para.add_run(title[:-1])
    _set_run_char_spacing(first, char_spacing, east_asia_hint=east_asia_hint)
    last = para.add_run(title[-1])
    if east_asia_hint:
        _set_run_east_asia_hint(last)
    return para


def _emit_page_break(anchor: Paragraph, doc) -> Paragraph:
    para = _new_paragraph_before(anchor, doc, S.S_NORMAL)
    para.add_run().add_break(WD_BREAK.PAGE)
    return para


def _clear_section_header_footer_refs(sectpr):
    for tag in ("w:headerReference", "w:footerReference"):
        for el in list(sectpr.findall(qn(tag))):
            sectpr.remove(el)


def _clear_section_page_number(sectpr):
    old = sectpr.find(qn("w:pgNumType"))
    if old is not None:
        sectpr.remove(old)


def _set_section_page_number(sectpr, fmt: Optional[str] = None, start: Optional[int] = None):
    _clear_section_page_number(sectpr)
    if fmt is None and start is None:
        return
    pg = OxmlElement("w:pgNumType")
    if fmt:
        pg.set(qn("w:fmt"), fmt)
    if start is not None:
        pg.set(qn("w:start"), str(start))
    sectpr.append(pg)


def _add_section_ref(sectpr, tag: str, ref_type: str, rel_id: Optional[str]):
    if not rel_id:
        return
    ref = OxmlElement(tag)
    ref.set(qn("w:type"), ref_type)
    ref.set(qn("r:id"), rel_id)
    sectpr.insert(0, ref)


def _rel_target_number(rel) -> int:
    m = re.search(r"(\d+)\.xml$", rel.target_ref or "")
    return int(m.group(1)) if m else 0


def _rel_blob(rel) -> str:
    try:
        return rel.target_part.blob.decode("utf-8", "ignore")
    except Exception:
        return ""


def _doc_style_names_by_id(doc) -> dict:
    return {
        getattr(style, "style_id", ""): getattr(style, "name", "")
        for style in doc.styles
        if getattr(style, "style_id", "")
    }


def _part_paragraph_style_names(part, style_names_by_id: dict) -> List[str]:
    names = []
    element = getattr(part, "element", None)
    if element is None:
        return names
    for paragraph in element.iter(qn("w:p")):
        ppr = paragraph.find(qn("w:pPr"))
        if ppr is None:
            continue
        pstyle = ppr.find(qn("w:pStyle"))
        if pstyle is None:
            continue
        style_id = pstyle.get(qn("w:val"))
        if style_id:
            names.append(style_names_by_id.get(style_id, style_id))
    return names


def _style_id_by_name(doc, style_name: str) -> Optional[str]:
    try:
        return doc.styles[style_name].style_id
    except KeyError:
        return None


def _set_part_paragraph_style(part, style_id: Optional[str]):
    if not style_id:
        return
    element = getattr(part, "element", None)
    if element is None:
        return
    for paragraph in element.iter(qn("w:p")):
        ppr = paragraph.find(qn("w:pPr"))
        if ppr is None:
            ppr = OxmlElement("w:pPr")
            paragraph.insert(0, ppr)
        pstyle = ppr.find(qn("w:pStyle"))
        if pstyle is None:
            pstyle = OxmlElement("w:pStyle")
            ppr.insert(0, pstyle)
        pstyle.set(qn("w:val"), style_id)
        # Let the applied standard header/footer style control alignment.
        # Some blueprint parts keep direct <w:jc w:val="right"/>, which
        # overrides the even-page header style until the style is reapplied in Word.
        for tag in ("w:jc",):
            old = ppr.find(qn(tag))
            if old is not None:
                ppr.remove(old)


def _choose_odd_even_rels(items, style_names_by_id: dict, odd_marker: str, even_marker: str):
    """Return (odd_rid, even_rid), preferring explicit standard odd/even styles."""
    if not items:
        return None, None
    odd = None
    even = None
    for item in items:
        names = _part_paragraph_style_names(item[1].target_part, style_names_by_id)
        if odd is None and any(odd_marker in name for name in names):
            odd = item
        if even is None and any(even_marker in name for name in names):
            even = item
    if odd is None:
        odd = items[0]
    if even is None:
        even = next((item for item in items if item is not odd), odd)
    return odd[0], even[0]


def _normalize_body_page_ref_styles(doc, refs: dict):
    style_map = {
        "header_default": _style_id_by_name(doc, S.S_HEADER_ODD),
        "header_even": _style_id_by_name(doc, S.S_HEADER_EVEN),
        "footer_default": _style_id_by_name(doc, S.S_FOOTER_ODD),
        "footer_even": _style_id_by_name(doc, S.S_FOOTER_EVEN),
    }
    applied_rids = set()
    for key, style_id in style_map.items():
        rid = refs.get(key)
        if not rid or rid not in doc.part.rels:
            continue
        if rid in applied_rids:
            continue
        _set_part_paragraph_style(doc.part.rels[rid].target_part, style_id)
        applied_rids.add(rid)


def _body_page_refs(doc):
    """查找封面蓝图中已打包的正文页眉/页脚关系。"""
    styleref_headers = []
    page_footers = []
    style_names_by_id = _doc_style_names_by_id(doc)
    for rid, rel in doc.part.rels.items():
        reltype = rel.reltype.rsplit("/", 1)[-1]
        blob = _rel_blob(rel)
        if reltype == "header" and "STYLEREF" in blob:
            styleref_headers.append((rid, rel))
        elif reltype == "footer" and "PAGE" in blob:
            page_footers.append((rid, rel))

    styleref_headers.sort(key=lambda item: _rel_target_number(item[1]))
    page_footers.sort(key=lambda item: _rel_target_number(item[1]))

    header_default, header_even = _choose_odd_even_rels(
        styleref_headers, style_names_by_id, "页眉奇数页", "页眉偶数页"
    )
    footer_default, footer_even = _choose_odd_even_rels(
        page_footers, style_names_by_id, "页脚奇数页", "页脚偶数页"
    )
    refs = {
        "header_default": header_default,
        "header_even": header_even,
        "footer_default": footer_default,
        "footer_even": footer_even,
    }
    _normalize_body_page_ref_styles(doc, refs)
    return refs


def _configure_section(sectpr, refs: Optional[dict] = None,
                       page_fmt: Optional[str] = None,
                       page_start: Optional[int] = None,
                       include_even: bool = False):
    type_el = sectpr.find(qn("w:type"))
    if type_el is None:
        type_el = OxmlElement("w:type")
        sectpr.insert(0, type_el)
    type_el.set(qn("w:val"), "nextPage")
    _clear_section_header_footer_refs(sectpr)
    if refs:
        _add_section_ref(sectpr, "w:headerReference", "default", refs.get("header_default"))
        _add_section_ref(sectpr, "w:footerReference", "default", refs.get("footer_default"))
        if include_even:
            _add_section_ref(sectpr, "w:headerReference", "even", refs.get("header_even"))
            _add_section_ref(sectpr, "w:footerReference", "even", refs.get("footer_even"))
    _set_section_page_number(sectpr, fmt=page_fmt, start=page_start)


def _new_section_break_before(anchor: Paragraph, doc, refs: Optional[dict] = None,
                              page_fmt: Optional[str] = None,
                              page_start: Optional[int] = None,
                              include_even: bool = False):
    """在 anchor 前插入下一页分节符，cover 后端按模板节结构切分各部分。"""
    new_p = OxmlElement("w:p")
    anchor._p.addprevious(new_p)
    para = Paragraph(new_p, anchor._parent)
    ppr = para._p.get_or_add_pPr()
    sectpr = copy.deepcopy(doc.sections[-1]._sectPr)
    _configure_section(
        sectpr,
        refs=refs,
        page_fmt=page_fmt,
        page_start=page_start,
        include_even=include_even,
    )
    ppr.append(sectpr)
    return para


def _configure_final_section(doc, refs: Optional[dict] = None,
                             page_start: Optional[int] = None,
                             include_even: bool = False):
    sectpr = doc.sections[-1]._sectPr
    _configure_section(
        sectpr,
        refs=refs,
        page_start=page_start,
        include_even=include_even,
    )


def _set_field(doc, style_name: str, text: str) -> bool:
    """把第一个套用 style_name 的段落文本替换为 text，保留首个 run 的格式。返回是否命中。"""
    for p in doc.paragraphs:
        if p.style is not None and p.style.name == style_name:
            _replace_text_keep_format(p, text)
            return True
    return False


def _set_field_or_placeholder(doc, style_name: str, text: str, patterns: List[str]) -> bool:
    if _set_field(doc, style_name, text):
        return True
    for p in doc.paragraphs:
        current = p.text.strip()
        if any(re.search(pattern, current) for pattern in patterns):
            _replace_text_keep_format(p, text)
            return True
    return False


def _replace_text_keep_format(p: Paragraph, text: str):
    runs = p.runs
    if runs:
        runs[0].text = text
        for extra in runs[1:]:
            extra._r.getparent().remove(extra._r)
    else:
        p.add_run(text)


def _find_para(doc, style_name: Optional[str] = None, text: Optional[str] = None,
               start: int = 0) -> Optional[int]:
    """在 doc.paragraphs 中查找匹配段落，返回索引。"""
    for i in range(start, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        if style_name is not None and (p.style is None or p.style.name != style_name):
            continue
        if text is not None and p.text.strip() != text:
            continue
        return i
    return None


def _next_sectpr_para(doc, start: int) -> Optional[int]:
    """从 start（含）起，找下一个承载 sectPr 的段落索引。"""
    for i in range(start, len(doc.paragraphs)):
        if _carries_sectpr(doc.paragraphs[i]):
            return i
    return None


def _prev_sectpr_para(doc, before: int) -> Optional[int]:
    """从 before 之前向前找最近的承载 sectPr 的段落索引。"""
    for i in range(before - 1, -1, -1):
        if _carries_sectpr(doc.paragraphs[i]):
            return i
    return None


def _find_first_appendix_mark(doc, before: Optional[int] = None) -> Optional[int]:
    """查找模板中第一个附录标识段，用于分离正文样例和附录样例。"""
    end = before if before is not None else len(doc.paragraphs)
    for i in range(0, end):
        p = doc.paragraphs[i]
        if p.style is not None and p.style.name == S.S_APPENDIX_MARK:
            return i
    return None


def _remove_between(doc, start_p: Paragraph, end_p: Paragraph):
    """删除 body 中 start_p 与 end_p 之间的所有元素（不含两端）。

    会顺带删除区间内的 sdt（占位控件）等非段落元素。
    """
    body = start_p._p.getparent()
    start_el = start_p._p
    end_el = end_p._p
    el = start_el.getnext()
    while el is not None and el is not end_el:
        nxt = el.getnext()
        body.remove(el)
        el = nxt


# --------------------------------------------------------------------------- #
# TOC 域
# --------------------------------------------------------------------------- #
def _make_toc_field(anchor: Paragraph, doc):
    """在 anchor 前插入一个 TOC 域段落（目录 1-3 级，含页码）。

    使用 w:fldSimple，配合 settings 的 updateFields，Word 打开时自动生成目次。
    """
    para = _new_paragraph_before(anchor, doc, S.S_NORMAL)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r' TOC \o "1-3" \h \z \u ')
    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = '右键此处选择"更新域"以生成目次。'
    run.append(t)
    fld.append(run)
    para._p.append(fld)
    return para


def _enable_update_fields(doc):
    """让 Word 打开时提示更新域（用于 TOC 页码）。"""
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        uf = OxmlElement("w:updateFields")
        uf.set(qn("w:val"), "true")
        settings.append(uf)


def _set_even_and_odd_headers(doc, enabled: bool):
    settings = doc.settings.element
    old = settings.find(qn("w:evenAndOddHeaders"))
    if enabled:
        if old is None:
            settings.append(OxmlElement("w:evenAndOddHeaders"))
    elif old is not None:
        settings.remove(old)


def _set_section_form_protection(sectpr, protected: bool):
    form_prot = sectpr.find(qn("w:formProt"))
    if protected:
        if form_prot is not None:
            sectpr.remove(form_prot)
        return
    if form_prot is None:
        form_prot = OxmlElement("w:formProt")
        sectpr.append(form_prot)
    form_prot.set(qn("w:val"), "0")


def _enable_cover_form_field_protection(doc):
    """Activate legacy form fields on the cover while keeping later sections editable."""
    settings = doc.settings.element
    protection = settings.find(qn("w:documentProtection"))
    if protection is None:
        protection = OxmlElement("w:documentProtection")
        settings.append(protection)
    for attr in list(protection.attrib):
        del protection.attrib[attr]
    protection.set(qn("w:edit"), "forms")
    protection.set(qn("w:enforcement"), "1")

    for idx, section in enumerate(doc.sections):
        _set_section_form_protection(section._sectPr, protected=(idx == 0))


def _disable_form_field_protection(doc):
    settings = doc.settings.element
    protection = settings.find(qn("w:documentProtection"))
    if protection is not None:
        for attr in list(protection.attrib):
            del protection.attrib[attr]
        protection.set(qn("w:edit"), "forms")
        protection.set(qn("w:enforcement"), "0")
    for section in doc.sections:
        _set_section_form_protection(section._sectPr, protected=False)


def _should_enable_cover_form_protection(
    meta: model.Meta,
    cover_form_protection: Optional[bool],
) -> bool:
    if cover_form_protection is not None:
        return bool(cover_form_protection)
    return bool(getattr(meta, "cover_form_protection", False))


def _legacy_dropdown_field_info(begin_run_el):
    fld = begin_run_el.find(qn("w:fldChar"))
    if fld is None:
        return None
    ffdata = fld.find(qn("w:ffData"))
    if ffdata is None:
        return None
    name_el = ffdata.find(qn("w:name"))
    ddlist = ffdata.find(qn("w:ddList"))
    if name_el is None or ddlist is None:
        return None
    return name_el.get(qn("w:val"), ""), ddlist


def _normalize_draft_version(value: str) -> str:
    text = (value or "").strip()
    aliases = {
        "工作组讨论稿": "（工作组讨论稿）",
        "征求意见稿": "（征求意见稿）",
        "送审讨论稿": "（送审讨论稿）",
        "送审稿": "（送审稿）",
        "报批稿": "（报批稿）",
    }
    return aliases.get(text, text)


def _set_legacy_dropdown_value(doc, field_name: str, value: str) -> bool:
    selected = _normalize_draft_version(value)
    if not selected:
        return False
    for run_el in doc.element.iter(qn("w:r")):
        fld = run_el.find(qn("w:fldChar"))
        if fld is None or fld.get(qn("w:fldCharType")) != "begin":
            continue
        info = _legacy_dropdown_field_info(run_el)
        if info is None:
            continue
        name, ddlist = info
        if name != field_name:
            continue

        entries = ddlist.findall(qn("w:listEntry"))
        values = [entry.get(qn("w:val"), "") for entry in entries]
        if selected not in values:
            entry = OxmlElement("w:listEntry")
            entry.set(qn("w:val"), selected)
            ddlist.append(entry)
            values.append(selected)
        idx = values.index(selected)
        result = ddlist.find(qn("w:result"))
        if result is None:
            result = OxmlElement("w:result")
            ddlist.insert(0, result)
        result.set(qn("w:val"), str(idx))
        return True
    return False


# --------------------------------------------------------------------------- #
# 各节构建
# --------------------------------------------------------------------------- #
def _apply_cover_fields(doc, meta: model.Meta, kind: Optional[str] = None):
    # ICS / CCS 表（封面第一个表格的两行第二列）
    if doc.tables:
        t = doc.tables[0]
        try:
            if meta.ics:
                _set_cell_text(t.rows[0].cells[1], meta.ics)
            if meta.ccs:
                _set_cell_text(t.rows[1].cells[1], meta.ccs)
        except IndexError:
            pass

    standard_type = meta.standard_type
    if kind == "national" and standard_type in ("", "国家标准", "中华人民共和国国家标准"):
        standard_type = "中华人民共和国国家标准"
    if standard_type:
        _set_field_or_placeholder(doc, S.S_COVER_TYPE, standard_type, [
            r"^团体标准$",
            r"^中华人民共和国国家标准$",
        ])
    if meta.number:
        _set_field_or_placeholder(doc, S.S_COVER_NUMBER, meta.number, [
            r"^T/XXX\s+XXXX",
            r"^GB/T\s+XXXXX",
            r"^GB\s+XXXXX",
        ])
    if meta.replaces:
        _set_field_or_placeholder(doc, S.S_COVER_REPLACES, "代替 %s" % meta.replaces, [
            r"^代替\s+",
        ])
    if meta.title:
        _set_field_or_placeholder(doc, S.S_COVER_NAME, meta.title, [
            r"^点击此处添加标准名称$",
        ])
    if meta.title_en:
        _set_field_or_placeholder(doc, S.S_COVER_NAME_EN, meta.title_en, [
            r"^点击此处添加标准名称的英文译名$",
        ])
    if meta.draft_version:
        _set_legacy_dropdown_value(doc, "下拉1", meta.draft_version)
    if meta.publish_date:
        _set_field_or_placeholder(doc, S.S_COVER_PUBLISH, "%s发布" % meta.publish_date, [
            r"^XXXX\s*-\s*XX\s*-\s*XX发布$",
        ])
    if meta.implement_date:
        _set_field_or_placeholder(doc, S.S_COVER_IMPLEMENT, "%s实施" % meta.implement_date, [
            r"^XXXX\s*-\s*XX\s*-\s*XX实施$",
        ])
    publisher_set = False
    if meta.publisher:
        publisher_set = _set_field_or_placeholder(doc, S.S_COVER_PUBLISHER, meta.publisher, [])
    return {"publisher": publisher_set}


_COVER_PLACEHOLDER_RE = re.compile(
    r"(\(?点击此处添加[^ \t\r\n<]*\)?)|"
    r"(点击此处添加[^ \t\r\n<]*)|"
    r"（本草案完成时间：）|"
    r"XXXX\s*-\s*XX\s*-\s*XX(?:发布|实施)?"
)


def _cleanup_placeholder_paragraph(p: Paragraph):
    text = p.text
    if not text or not _COVER_PLACEHOLDER_RE.search(text):
        return
    cleaned = _COVER_PLACEHOLDER_RE.sub("", text).strip()
    _replace_text_keep_format(p, cleaned)


def _cleanup_cover_placeholders(doc):
    for p in doc.paragraphs:
        _cleanup_placeholder_paragraph(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _cleanup_placeholder_paragraph(p)


def _find_first_section_break_paragraph(doc) -> Optional[Paragraph]:
    for p in doc.paragraphs:
        if _carries_sectpr(p):
            return p
    return None


def _ensure_cover_publisher(doc, publisher: str, cover_info: Optional[dict], kind: str = "group"):
    if not publisher or (cover_info or {}).get("publisher"):
        return
    # 国家标准封面蓝图底部发布单位为打包图片；再插入文本段会与图片重叠。
    if kind == "national":
        return
    anchor = _find_first_section_break_paragraph(doc)
    if anchor is None:
        para = doc.add_paragraph()
    else:
        para = _new_paragraph_before(anchor, doc, S.S_COVER_PUBLISHER)
    if not para.runs:
        para.add_run(publisher)
    else:
        _replace_text_keep_format(para, publisher)


def _set_cell_text(cell, text):
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = text
        for extra in p.runs[1:]:
            extra._r.getparent().remove(extra._r)
    else:
        p.add_run(text)


def _build_toc(doc):
    title_idx = _find_para(doc, style_name=S.S_TOC_TITLE, text="目次")
    if title_idx is None:
        return
    title_p = doc.paragraphs[title_idx]
    term_idx = _next_sectpr_para(doc, title_idx + 1)
    if term_idx is None:
        return
    term_p = doc.paragraphs[term_idx]
    _remove_between(doc, title_p, term_p)
    _make_toc_field(term_p, doc)


def _emit_foreword_content(anchor: Paragraph, doc, meta: model.Meta):
    fw = meta.foreword

    def add(text):
        if text:
            _new_paragraph_before(anchor, doc, S.S_PARA, text=text)

    def add_extra_note(note):
        if isinstance(note, (list, tuple)):
            for item in note:
                text = str(item).strip()
                if text.startswith("- "):
                    text = text[2:].strip()
                if text:
                    _new_numbered_style_paragraph(anchor, doc, S.S_LIST_DASH, text=text)
            return
        add(str(note).strip())

    add(bp.FOREWORD_FIRST)
    if fw.multipart_note:
        add(fw.multipart_note)
    if fw.replace_changes:
        # 技术变化导语 + 破折号列项
        if meta.replaces:
            lead = ("本文件代替%s，与%s相比，除结构调整和编辑性改动外，"
                    "主要技术变化如下：" % (meta.replaces, meta.replaces))
        else:
            lead = "与上一版相比，除结构调整和编辑性改动外，主要技术变化如下："
        add(lead)
        for ch in fw.replace_changes:
            _new_numbered_style_paragraph(anchor, doc, S.S_LIST_DASH, text=ch)
    if fw.patent_note:
        add(bp.PATENT_NOTE)
    # 提出 / 归口
    if fw.proposer and (not fw.owner or fw.owner == fw.proposer):
        add(bp.PROPOSE_OWNER_SAME.format(org=fw.proposer))
    else:
        if fw.proposer:
            add(bp.PROPOSE.format(org=fw.proposer))
        if fw.owner:
            add(bp.OWNER.format(org=fw.owner))
    if fw.draft_orgs:
        add(bp.DRAFT_ORGS.format(orgs="、".join(fw.draft_orgs)))
    if fw.drafters:
        add(bp.DRAFTERS.format(people="、".join(fw.drafters)))
    if fw.history:
        add(bp.HISTORY_LEAD)
        add(fw.history)
    for note in fw.extra_notes:
        add_extra_note(note)


def _build_foreword(doc, meta: model.Meta):
    idx = _find_para(doc, style_name=S.S_PREFACE_TITLE, text="前言")
    if idx is None:
        return
    title_p = doc.paragraphs[idx]
    term_idx = _next_sectpr_para(doc, idx + 1)
    if term_idx is None:
        return
    term_p = doc.paragraphs[term_idx]
    _remove_between(doc, title_p, term_p)
    _emit_foreword_content(term_p, doc, meta)


def _emit_introduction_content(anchor: Paragraph, doc, meta: model.Meta):
    for line in meta.introduction.splitlines():
        if line.strip():
            _new_paragraph_before(anchor, doc, S.S_PARA, text=line.strip())


def _build_introduction(doc, meta: model.Meta):
    idx = _find_para(doc, style_name=S.S_PREFACE_TITLE, text="引言")
    if idx is None:
        return
    title_p = doc.paragraphs[idx]
    term_idx = _next_sectpr_para(doc, idx + 1)
    if term_idx is None:
        return
    term_p = doc.paragraphs[term_idx]

    if meta.introduction.strip():
        _remove_between(doc, title_p, term_p)
        _emit_introduction_content(term_p, doc, meta)
    else:
        # 无引言：删除标题 + 区间内容 + 本节终止段（含其 sectPr），整节合并入正文
        _remove_between(doc, title_p, term_p)
        body = title_p._p.getparent()
        body.remove(title_p._p)
        body.remove(term_p._p)


_TERM_SPLIT_RE = re.compile(r"^(.*?)[\s　]{2,}(.+)$")
# 规范性引用文件条目："标准号  标准名称"（标准号如 GB 5749 / GB/T 5750.3 / DZ/T 0225）
_NORMREF_RE = re.compile(r"^\s*([A-Z][A-Z/]*\s+\d[\w.\-—–]*)(?:\s{2,}|　+)(.+)$")


def _emit_normative_ref(anchor: Paragraph, doc, spans):
    """规范性引用文件条目：把标准号加书签，便于交叉引用插入"GB 5749"。"""
    text = "".join(s.text for s in spans).strip()
    m = _NORMREF_RE.match(text)
    if not m:
        _new_paragraph_before(anchor, doc, S.S_PARA, spans=spans)
        return
    stdno, name = m.group(1).strip(), m.group(2).strip()
    para = _new_paragraph_before(anchor, doc, S.S_PARA)
    bid = _COUNTER.bm
    _COUNTER.bm += 1
    _bookmark_start(para, _bm_name(stdno), bid)
    para.add_run(stdno)
    _bookmark_end(para, bid)
    para.add_run("  " + name)


def _emit_source(anchor: Paragraph, doc, source: model.Source):
    para = _new_paragraph_before(anchor, doc, S.S_PARA)
    run = para.add_run(source.text)
    try:
        run.style = doc.styles[S.S_SOURCE]
    except KeyError:
        pass


def _emit_term(anchor: Paragraph, doc, spans):
    """术语条目：编号(3.1)单独成行，下一行加粗中文术语 + 英文对应词。

    自动编号接入 numId=2 ilvl=2（术语为章 3 下的一级条）。
    """
    text = "".join(s.text for s in spans).strip()
    m = _TERM_SPLIT_RE.match(text)
    if m:
        cn, en = m.group(1).strip(), m.group(2).strip()
    else:
        cn, en = text, ""
    para = _new_paragraph_before(anchor, doc, S.S_TERM_1)
    _set_numbering(para, S.NUM_BODY, 2)
    para.add_run().add_break()          # 让自动编号 3.1 单独占一行
    r = para.add_run(cn)
    r.bold = True
    if en:
        para.add_run("　")
        er = para.add_run(en)
        er.bold = True
    return para


def _emit_body_block(anchor: Paragraph, doc, blk, in_terms=False, in_normrefs=False,
                     appendix_letter=None, list_num_id: Optional[int] = None):
    """把一个正文块插入到 anchor 之前。"""
    if isinstance(blk, model.Heading):
        if in_terms and blk.level == 2:
            _emit_term(anchor, doc, blk.spans)
            return
        style = S.HEADING_STYLE_BY_LEVEL.get(blk.level, S.S_PARA)
        _new_paragraph_before(anchor, doc, style, spans=blk.spans)
    elif isinstance(blk, model.Paragraph):
        if in_normrefs:
            _emit_normative_ref(anchor, doc, blk.spans)
        else:
            _new_paragraph_before(anchor, doc, S.S_PARA, spans=blk.spans)
    elif isinstance(blk, model.UntitledClause):
        # 无标题条：套用"X级无标题"样式 + 接入正文多级列表(numId=2)自动编号，
        # Markdown 只声明层级；ilvl = 层级（3 -> 渲染为 "4.2.1"）。
        seg = blk.segments
        style = S.UNTITLED_STYLE_BY_SEGMENTS.get(seg, S.S_PARA)
        para = _new_paragraph_before(anchor, doc, style, spans=list(blk.spans))
        if seg in S.UNTITLED_STYLE_BY_SEGMENTS:
            _set_numbering(para, S.NUM_BODY, seg)
    elif isinstance(blk, model.Note):
        style = S.S_NOTE_X if blk.index else S.S_NOTE
        _new_paragraph_before(anchor, doc, style, spans=blk.spans)
    elif isinstance(blk, model.Example):
        style = S.S_EXAMPLE
        _new_paragraph_before(anchor, doc, style, spans=blk.spans)
    elif isinstance(blk, model.ExampleContent):
        _new_paragraph_before(anchor, doc, S.S_EXAMPLE_CONTENT, spans=blk.spans)
    elif isinstance(blk, model.PageBreak):
        _emit_page_break(anchor, doc)
    elif isinstance(blk, model.Source):
        _emit_source(anchor, doc, blk)
    elif isinstance(blk, model.ListBlock):
        if blk.ordered:
            style = S.ORDERED_LIST_STYLE_BY_LEVEL.get(blk.level, S.S_LIST_NUMBER_3)
        else:
            style = S.S_LIST_DASH
        for it in blk.items:
            _new_numbered_style_paragraph(
                anchor, doc, style, spans=it.spans,
                num_id_override=list_num_id if blk.ordered else None,
            )
    elif isinstance(blk, model.TableModel):
        _emit_table(anchor, doc, blk, appendix_letter=appendix_letter)
    elif isinstance(blk, model.Figure):
        _emit_figure(anchor, doc, blk, appendix_letter=appendix_letter)
    elif isinstance(blk, model.Formula):
        fstyle = S.S_FORMULA_APPENDIX if appendix_letter else S.S_FORMULA
        _emit_formula(anchor, doc, blk, style=fstyle, appendix_letter=appendix_letter)


def _emit_body_standard_title(anchor: Paragraph, doc, meta: model.Meta):
    title = (meta.title or "").strip()
    if not title:
        return
    para = _new_paragraph_before(anchor, doc, S.S_BODY_STANDARD_NAME)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, line in enumerate([x.strip() for x in title.splitlines() if x.strip()]):
        if i > 0:
            para.add_run().add_break()
        run = para.add_run(line)
        run.bold = True


def _emit_formula(anchor: Paragraph, doc, formula, style=None, appendix_letter=None):
    """公式段：用制表位——[Tab]公式[Tab]（序号）。

    模板"标准文件_正文公式"已配好：居中制表位让公式居中，右制表位(带点引导)推序号。
    序号括在"（""）"内，括号为纯文本；书签仅围住不带括号的编号文字。
    SEQ 域使用中文前缀"公式"，使公式出现在 Word 交叉引用"公式"域。
    """
    from . import mathconv
    _emit_formula_caption_anchor(anchor, doc, formula, appendix_letter=appendix_letter)
    para = _new_paragraph_before(anchor, doc, style or S.S_FORMULA)
    # 居中制表位：公式居中
    para.add_run("\t")
    # 原生公式（OMML）
    omath = mathconv.latex_to_omml(formula.latex)
    if omath is not None:
        para._p.append(omath)
    else:
        para.add_run(formula.latex)
    para.add_run("\t")
    _emit_formula_number(para, formula, appendix_letter=appendix_letter)


# 表/图编号计数器（模块级在 build() 内重置）
class _Counter:
    def __init__(self):
        self.table = 0
        self.figure = 0
        self.bm = 1000
        self.seq_scope_counts = {}


_COUNTER = _Counter()


def _needs_seq_reset(ref_type: str, appendix_letter=None) -> bool:
    scope = appendix_letter or "body"
    key = (ref_type, scope)
    count = _COUNTER.seq_scope_counts.get(key, 0)
    _COUNTER.seq_scope_counts[key] = count + 1
    return count == 0


def _add_bookmark_ends_after(run, bids):
    """按 bids 给出的最终顺序，把多个 bookmarkEnd 插到同一个 run 后。"""
    for bid in reversed([b for b in bids if b is not None]):
        _bookmark_end_after(run, bid)


def _emit_visible_caption(para: Paragraph, ref_type: str, seq_name: str,
                          label_text: str, anchor_id: str, title: str,
                          appendix_letter=None):
    """输出 `表1　标题` / `图A.1　标题`，并建立 num/label/full 书签。"""
    label_run = para.add_run(label_text)
    bm_full_id = bm_label_id = bm_num_id = bm_text_id = None
    if anchor_id:
        bm_full_id = _next_bm_id()
        bm_label_id = _next_bm_id()
        bm_num_id = _next_bm_id()
        bm_text_id = _next_bm_id() if title else None
        _bookmark_start_before(label_run, _native_ref_name(ref_type, anchor_id, "full"), bm_full_id)
        _bookmark_start_before(label_run, _native_ref_name(ref_type, anchor_id, "label"), bm_label_id)

    prefix_run = None
    if appendix_letter:
        prefix_run = para.add_run(appendix_letter + ".")
    result_run, seq_end_run = _add_seq(para, seq_name, reset=_needs_seq_reset(ref_type, appendix_letter))

    if anchor_id:
        _bookmark_start_before(prefix_run or result_run, _native_ref_name(ref_type, anchor_id, "num"), bm_num_id)

    title_run = None
    if title:
        para.add_run("　")
        title_run = para.add_run(title)
        if anchor_id and bm_text_id is not None:
            _bookmark_start_before(title_run, _native_ref_name(ref_type, anchor_id, "text"), bm_text_id)

    if anchor_id:
        if title_run is not None:
            _add_bookmark_ends_after(seq_end_run, [bm_num_id, bm_label_id])
            if bm_text_id is not None:
                _bookmark_end_after(title_run, bm_text_id)
            _bookmark_end_after(title_run, bm_full_id)
        else:
            _add_bookmark_ends_after(seq_end_run, [bm_num_id, bm_label_id, bm_full_id])


def _emit_formula_number(para: Paragraph, formula: model.Formula, appendix_letter=None):
    """输出公式右侧编号 `（1）` / `（A.1）`，显示值引用隐藏公式题注。"""
    para.add_run("（")
    placeholder = "1"
    if appendix_letter:
        placeholder = appendix_letter + ".1"
    if formula.anchor_id:
        _add_ref_bookmark(
            para,
            _native_ref_name("eq", formula.anchor_id, "num"),
            display_text=placeholder,
            charformat=True,
        )
    else:
        para.add_run(placeholder)
    para.add_run("）")


def _emit_formula_caption_anchor(anchor: Paragraph, doc, formula: model.Formula, appendix_letter=None):
    """插入一个隐藏公式题注段，供 Word 交叉引用列表和 REF 域使用。"""
    if not formula.anchor_id:
        return
    cap = _new_paragraph_before(anchor, doc, S.S_NORMAL)
    try:
        cap.style = doc.styles["Caption"]
    except KeyError:
        pass
    _make_paragraph_hidden(cap)

    label_run = cap.add_run("公式")
    _make_run_hidden(label_run._r)
    bm_label_id = _next_bm_id()
    bm_num_id = _next_bm_id()
    _bookmark_start_before(label_run, _native_ref_name("eq", formula.anchor_id, "label"), bm_label_id)
    space_run = cap.add_run(" ")
    _make_run_hidden(space_run._r)
    prefix_run = None
    if appendix_letter:
        prefix_run = cap.add_run(appendix_letter + ".")
        _make_run_hidden(prefix_run._r)
    result_run, seq_end_run = _add_seq(
        cap,
        SEQ_EQUATION,
        reset=_needs_seq_reset("eq", appendix_letter),
        hidden=True,
    )
    _bookmark_start_before(prefix_run or result_run, _native_ref_name("eq", formula.anchor_id, "num"), bm_num_id)
    _add_bookmark_ends_after(seq_end_run, [bm_num_id, bm_label_id])


# 不按行数预拆续表。是否跨页只有 Word/LibreOffice 完成分页后才知道；
# 生成阶段误拆会产生同页“续表”，因此默认只生成一个真实表格。
_TABLE_SPLIT_THRESHOLD = None
_TABLE_SPLIT_CHUNK_SIZE = None


def _set_cell_vertical_center(cell):
    tcpr = cell._tc.get_or_add_tcPr()
    old = tcpr.find(qn("w:vAlign"))
    if old is not None:
        tcpr.remove(old)
    valign = OxmlElement("w:vAlign")
    valign.set(qn("w:val"), "center")
    tcpr.append(valign)


def _set_row_repeat_header(row):
    trpr = row._tr.get_or_add_trPr()
    if trpr.find(qn("w:tblHeader")) is None:
        hdr = OxmlElement("w:tblHeader")
        hdr.set(qn("w:val"), "true")
        trpr.append(hdr)


def _cell_is_long_text(text: str, colspan: int = 1) -> bool:
    text = (text or "").strip()
    return colspan > 1 or len(text) > 25 or "。" in text or "；" in text or "：" in text


def _parts_from_text(text: str) -> List[model.TableCellPart]:
    return [model.TableCellPart("text", text or "")]


def _emit_table_cell_parts(paragraph: Paragraph, parts: List[model.TableCellPart]):
    from . import mathconv
    if not parts:
        return
    for part in parts:
        if part.kind == "ref":
            _add_typed_ref(paragraph, model.RefSpan(
                text=part.text,
                ref_type=part.ref_type,
                target=part.target,
                mode=part.mode,
            ))
            continue
        if part.kind == "formula":
            omath = mathconv.latex_to_omml(part.text)
            if omath is not None:
                paragraph._p.append(omath)
            else:
                paragraph.add_run(part.text)
            continue
        for i, piece in enumerate((part.text or "").split("\n")):
            if i > 0:
                paragraph.add_run().add_break()
            if piece:
                paragraph.add_run(piece)


def _emit_table_continuation_caption(anchor: Paragraph, doc, tbl: model.TableModel,
                                     appendix_letter=None):
    style = "标准文件_表格续"
    cap = _new_paragraph_before(anchor, doc, style)
    _set_numbering(cap, 0, 0)
    cap.paragraph_format.page_break_before = True
    if tbl.anchor_id:
        _add_ref_bookmark(
            cap,
            _native_ref_name("tbl", tbl.anchor_id, "label"),
            display_text="表?",
        )
    else:
        cap.add_run("表")
    title = (tbl.caption or "").strip()
    if title:
        cap.add_run("　")
        cap.add_run(title)
    cap.add_run("（续）")


def _emit_table_part(anchor: Paragraph, doc, tbl: model.TableModel, rows, row_parts, row_colspans,
                     appendix_letter=None):
    """输出一个表格片段。续表复用原表头，不新增 SEQ。"""

    def row_width(row, spans):
        if spans:
            return sum(spans[:len(row)])
        return len(row)

    widths = []
    if tbl.header:
        widths.append(row_width(tbl.header, tbl.header_colspans))
    widths.extend(row_width(row, row_colspans[i] if i < len(row_colspans) else [])
                  for i, row in enumerate(rows))
    ncols = max(widths) if widths else 1
    table = doc.add_table(rows=0, cols=ncols)
    try:
        table.style = doc.styles[S.S_TABLE_GRID]
    except KeyError:
        pass

    all_rows = ([tbl.header] if tbl.header else []) + rows
    header_parts = tbl.header_parts or [_parts_from_text(text) for text in tbl.header]
    all_parts = ([header_parts] if tbl.header else []) + row_parts
    all_spans = ([tbl.header_colspans or [1 for _ in tbl.header]] if tbl.header else [])
    all_spans += [
        row_colspans[i] if i < len(row_colspans) and row_colspans[i]
        else [1 for _ in row]
        for i, row in enumerate(rows)
    ]

    for r_i, row in enumerate(all_rows):
        word_row = table.add_row()
        if r_i == 0 and tbl.header:
            _set_row_repeat_header(word_row)
        cells = word_row.cells
        spans = all_spans[r_i] if r_i < len(all_spans) else [1 for _ in row]
        parts_row = all_parts[r_i] if r_i < len(all_parts) else []
        c_pos = 0
        for c_i, txt in enumerate(row):
            if c_pos >= ncols:
                break
            colspan = spans[c_i] if c_i < len(spans) else 1
            colspan = max(1, min(colspan, ncols - c_pos))
            cell = cells[c_pos]
            if colspan > 1:
                cell = cell.merge(cells[c_pos + colspan - 1])
            _set_cell_vertical_center(cell)
            cp = cell.paragraphs[0]
            try:
                cp.style = doc.styles[S.S_TABLE_CELL]
            except KeyError:
                pass
            cp.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT
                if _cell_is_long_text(txt, colspan)
                else WD_ALIGN_PARAGRAPH.CENTER
            )
            parts = parts_row[c_i] if c_i < len(parts_row) else _parts_from_text(txt)
            _emit_table_cell_parts(cp, parts)
            c_pos += colspan
    anchor._p.addprevious(table._tbl)


def _emit_table(anchor: Paragraph, doc, tbl: model.TableModel, appendix_letter=None):
    """表格：标题显式输出 `表N　标题`，模板样式只负责排版。"""
    style = S.S_APPENDIX_TABLE_CAPTION if appendix_letter else S.S_TABLE_CAPTION
    cap = _new_paragraph_before(anchor, doc, style)
    _set_numbering(cap, 0, 0)
    _emit_visible_caption(
        cap, "tbl", SEQ_TABLE, "表", tbl.anchor_id, tbl.caption or "",
        appendix_letter=appendix_letter,
    )

    row_parts = tbl.row_parts or [
        [_parts_from_text(cell) for cell in row]
        for row in tbl.rows
    ]
    row_colspans = tbl.row_colspans or [
        [1 for _ in row]
        for row in tbl.rows
    ]
    if _TABLE_SPLIT_THRESHOLD is not None and len(tbl.rows) > _TABLE_SPLIT_THRESHOLD:
        start = 0
        first = True
        while start < len(tbl.rows):
            end = min(start + _TABLE_SPLIT_CHUNK_SIZE, len(tbl.rows))
            if not first:
                _emit_table_continuation_caption(anchor, doc, tbl, appendix_letter=appendix_letter)
            _emit_table_part(
                anchor,
                doc,
                tbl,
                tbl.rows[start:end],
                row_parts[start:end],
                row_colspans[start:end],
                appendix_letter=appendix_letter,
            )
            first = False
            start = end
        return

    _emit_table_part(anchor, doc, tbl, tbl.rows, row_parts, row_colspans, appendix_letter=appendix_letter)


def _emit_figure(anchor: Paragraph, doc, fig: model.Figure, appendix_letter=None):
    img_p = _new_paragraph_before(anchor, doc, S.S_NORMAL)
    if fig.path and os.path.isfile(fig.path):
        try:
            img_p.add_run().add_picture(fig.path)
        except Exception:
            img_p.add_run("[图片无法加载：%s]" % fig.path)
    else:
        img_p.add_run("[缺少图片：%s]" % fig.path)
    # 图题：标题显式输出 `图N　标题`，模板样式只负责排版。
    style = S.S_APPENDIX_FIGURE_CAPTION if appendix_letter else S.S_FIGURE_CAPTION
    cap = _new_paragraph_before(anchor, doc, style)
    _set_numbering(cap, 0, 0)
    _emit_visible_caption(
        cap, "fig", SEQ_FIGURE, "图", fig.anchor_id, fig.caption or "",
        appendix_letter=appendix_letter,
    )


def _emit_body_blocks(anchor: Paragraph, doc, blocks: List[object]):
    blocks = _inject_chapter_boilerplate(blocks)
    in_terms = False
    in_normrefs = False
    ordered_list_num_id = None
    for blk in blocks:
        if isinstance(blk, model.Heading) and blk.level == 1:
            t = blk.text.strip()
            in_terms = "术语" in t and "定义" in t
            in_normrefs = "规范性引用文件" in t
        if isinstance(blk, model.ListBlock) and blk.ordered:
            style = S.ORDERED_LIST_STYLE_BY_LEVEL.get(blk.level, S.S_LIST_NUMBER_3)
            if blk.level == 1 or ordered_list_num_id is None:
                ordered_list_num_id = _new_numbering_instance_from_style(doc, style)
            _emit_body_block(
                anchor, doc, blk, in_terms=in_terms, in_normrefs=in_normrefs,
                list_num_id=ordered_list_num_id,
            )
        else:
            _emit_body_block(anchor, doc, blk, in_terms=in_terms, in_normrefs=in_normrefs)
            ordered_list_num_id = None


def _build_body(doc, sdoc: model.StandardDoc):
    """正文：定位正文节（引言/前言 sectPr 之后到 body-sectPr），清空样例后插入。"""
    ref_idx = _find_para(doc, style_name=S.S_REF_TITLE, text="参考文献")
    if ref_idx is None:
        return
    # 国家模板在参考文献前可能带有示例附录节；正文边界应取首个附录前的分节符。
    appendix_idx = _find_first_appendix_mark(doc, before=ref_idx)
    boundary_idx = appendix_idx if appendix_idx is not None else ref_idx
    term_idx = _prev_sectpr_para(doc, boundary_idx)
    if term_idx is None:
        return
    term_p = doc.paragraphs[term_idx]

    # 正文节起点：term 之前最近的另一个 sectPr 段（即引言/前言节终止段）
    start_idx = _prev_sectpr_para(doc, term_idx)
    start_p = doc.paragraphs[start_idx] if start_idx is not None else None

    # 清除正文节内既有样例（起点之后到终止段之间）
    if start_p is not None:
        _remove_between(doc, start_p, term_p)
    else:
        # 没有上游 sectPr，则清除文档开头到 term 之间——不应发生，保守跳过
        pass

    _emit_body_standard_title(term_p, doc, sdoc.meta)
    _emit_body_blocks(term_p, doc, sdoc.body)


def _inject_chapter_boilerplate(body: List[object]) -> List[object]:
    """为 规范性引用文件 / 术语和定义 / 符号和缩略语 空章补默认导语。"""
    out: List[object] = []
    n = len(body)
    for i, blk in enumerate(body):
        out.append(blk)
        if isinstance(blk, model.Heading) and blk.level == 1:
            title = blk.text.strip()
            # 判断该章后是否紧跟另一个一级章（即本章无正文内容）
            has_content = False
            for j in range(i + 1, n):
                nb = body[j]
                if isinstance(nb, model.Heading) and nb.level == 1:
                    break
                has_content = True
                break
            if not has_content:
                if title == S.CH_NORMATIVE_REF:
                    out.append(model.Paragraph(spans=[model.Span(bp.NORMATIVE_REF_NONE)]))
                elif title == S.CH_TERMS:
                    out.append(model.Paragraph(spans=[model.Span(bp.TERMS_NONE)]))
    return out


def _emit_appendices(anchor: Paragraph, doc, appendices: List[model.Appendix],
                     page_break_before: bool = True, section_before_each: bool = False,
                     start_index: int = 0):
    for ai, appx in enumerate(appendices):
        if section_before_each:
            _new_section_break_before(anchor, doc)
        letter = chr(ord("A") + start_index + ai)
        # 附录标题块：同一段三行，避免把附录标识、性质、标题拆成段落符。
        nature = bp.APPENDIX_NORMATIVE if appx.nature == "normative" else bp.APPENDIX_INFORMATIVE
        head = _new_paragraph_before(anchor, doc, S.S_APPENDIX_MARK)  # 自动"附录A"
        head.paragraph_format.page_break_before = page_break_before
        head.add_run().add_break()
        nature_run = head.add_run(nature)
        _apply_style_run_properties(doc, nature_run, S.S_APPENDIX_NATURE)
        head.add_run().add_break()
        _add_styled_runs(head, doc, S.S_APPENDIX_TITLE, appx.title_spans)
        ordered_list_num_id = None
        for blk in appx.blocks:
            if isinstance(blk, model.Heading):
                style = S.APPENDIX_CLAUSE_STYLE_BY_LEVEL.get(blk.level, S.S_PARA)
                _new_paragraph_before(anchor, doc, style, spans=blk.spans)
                ordered_list_num_id = None
            elif isinstance(blk, model.ListBlock) and blk.ordered:
                style = S.ORDERED_LIST_STYLE_BY_LEVEL.get(blk.level, S.S_LIST_NUMBER_3)
                if blk.level == 1 or ordered_list_num_id is None:
                    ordered_list_num_id = _new_numbering_instance_from_style(doc, style)
                _emit_body_block(
                    anchor, doc, blk,
                    appendix_letter=letter,
                    list_num_id=ordered_list_num_id,
                )
            else:
                _emit_body_block(anchor, doc, blk, appendix_letter=letter)
                ordered_list_num_id = None


def _clear_template_appendix_placeholders(doc):
    """删除完整模板自带的示例附录节，避免国家模板残留附录A/B样例。"""
    ref_idx = _find_para(doc, style_name=S.S_REF_TITLE, text="参考文献")
    if ref_idx is None:
        return
    appendix_idx = _find_first_appendix_mark(doc, before=ref_idx)
    if appendix_idx is None:
        return
    start_idx = _prev_sectpr_para(doc, appendix_idx)
    if start_idx is None:
        return
    _remove_between(doc, doc.paragraphs[start_idx], doc.paragraphs[ref_idx])


def _build_appendices(doc, sdoc: model.StandardDoc):
    _clear_template_appendix_placeholders(doc)
    ref_idx = _find_para(doc, style_name=S.S_REF_TITLE, text="参考文献")
    if ref_idx is None:
        return
    if not sdoc.appendices:
        return
    # 附录应位于正文节内、参考文献分节符之前：取参考文献标题前最近的承载 sectPr 的段落
    term_idx = _prev_sectpr_para(doc, ref_idx)
    anchor = doc.paragraphs[term_idx] if term_idx is not None else doc.paragraphs[ref_idx]
    _emit_appendices(anchor, doc, sdoc.appendices)


def _move_index_end_line_before_section_break(title_p: Paragraph):
    """把模板索引节末尾的居中线图移到上一节末尾，避免单独占尾页。"""
    body = title_p._p.getparent()
    section_break = title_p._p.getprevious()
    while section_break is not None:
        if section_break.tag == qn("w:p"):
            ppr = section_break.find(qn("w:pPr"))
            if ppr is not None and ppr.find(qn("w:sectPr")) is not None:
                break
        section_break = section_break.getprevious()
    if section_break is None:
        return

    el = title_p._p.getnext()
    while el is not None:
        nxt = el.getnext()
        if el.tag == qn("w:p"):
            ppr = el.find(qn("w:pPr"))
            if ppr is not None and ppr.find(qn("w:sectPr")) is not None:
                break
            txt = "".join(t.text or "" for t in el.findall(".//" + qn("w:t")))
            has_draw = el.findall(".//" + qn("w:drawing"))
            if has_draw and not txt.strip():
                section_break.addprevious(el)
        el = nxt
    body.remove(section_break)


def _find_index_end_line(title_p: Paragraph):
    """查找索引节末尾的标准文档结束线段。"""
    el = title_p._p.getnext()
    while el is not None:
        if el.tag == qn("w:p"):
            txt = "".join(t.text or "" for t in el.findall(".//" + qn("w:t")))
            has_draw = el.findall(".//" + qn("w:drawing"))
            if has_draw and not txt.strip():
                return el
        el = el.getnext()
    return None


def _clear_index_placeholders(title_p: Paragraph, end_line_el):
    """删除索引标题与结束线之间的模板占位段。"""
    body = title_p._p.getparent()
    el = title_p._p.getnext()
    while el is not None and el is not end_line_el:
        nxt = el.getnext()
        if el.tag == qn("w:p"):
            body.remove(el)
        el = nxt


def _emit_index_item(anchor: Paragraph, doc, item: model.IndexItem):
    para = _new_paragraph_before(anchor, doc, S.S_INDEX_ITEM)
    para.add_run(item.term)
    para.add_run("\t")
    para.add_run(item.targets)


def _build_references(doc, sdoc: model.StandardDoc):
    idx = _find_para(doc, style_name=S.S_REF_TITLE, text="参考文献")
    if idx is None:
        return
    title_p = doc.paragraphs[idx]
    term_idx = _next_sectpr_para(doc, idx + 1)
    if term_idx is None:
        return
    term_p = doc.paragraphs[term_idx]
    _remove_between(doc, title_p, term_p)
    if not sdoc.references:
        body = title_p._p.getparent()
        prev = title_p._p.getprevious()
        if prev is not None and prev.tag == qn("w:p"):
            ppr = prev.find(qn("w:pPr"))
            if ppr is not None and ppr.find(qn("w:sectPr")) is not None:
                body.remove(prev)
        body.remove(title_p._p)
        return
    for item in sdoc.references:
        _new_paragraph_before(term_p, doc, S.S_REF_ITEM, text=item)


def _build_index(doc, sdoc: model.StandardDoc):
    """生成或删除索引节。"""
    idx = _find_para(doc, style_name=S.S_INDEX_TITLE, text="索引")
    if idx is None:
        return
    title_p = doc.paragraphs[idx]
    if sdoc.index_groups:
        end_line_el = _find_index_end_line(title_p)
        if end_line_el is None:
            return
        _clear_index_placeholders(title_p, end_line_el)
        anchor = Paragraph(end_line_el, title_p._parent)
        for group in sdoc.index_groups:
            _new_paragraph_before(anchor, doc, S.S_INDEX_LETTER, text=group.letter)
            for item in group.items:
                _emit_index_item(anchor, doc, item)
        return

    _move_index_end_line_before_section_break(title_p)
    body = title_p._p.getparent()
    # 删除标题之后、直到下一个承载 sectPr 的段或文末的空段
    el = title_p._p.getnext()
    while el is not None:
        nxt = el.getnext()
        if el.tag == qn("w:p"):
            ppr = el.find(qn("w:pPr"))
            if ppr is not None and ppr.find(qn("w:sectPr")) is not None:
                break
            txt = "".join(t.text or "" for t in el.findall(".//" + qn("w:t")))
            has_draw = el.findall(".//" + qn("w:drawing"))
            if txt.strip() or has_draw:
                # 保留有图/有文字的段（如背景图）
                el = nxt
                continue
            body.remove(el)
        el = nxt
    if el is not None and el.tag == qn("w:p"):
        ppr = el.find(qn("w:pPr"))
        if ppr is not None and ppr.find(qn("w:sectPr")) is not None:
            body.remove(el)
    body.remove(title_p._p)


def _remove_paragraph(para: Paragraph):
    parent = para._p.getparent()
    parent.remove(para._p)


def _emit_cover_toc(anchor: Paragraph, doc, kind: str):
    _emit_section_title_before(anchor, doc, S.S_TOC_TITLE, "目次", kind=kind)
    _make_toc_field(anchor, doc)


def _emit_cover_foreword(anchor: Paragraph, doc, meta: model.Meta, kind: str):
    _emit_section_title_before(anchor, doc, S.S_PREFACE_TITLE, "前言", kind=kind)
    _emit_foreword_content(anchor, doc, meta)


def _emit_cover_introduction(anchor: Paragraph, doc, meta: model.Meta, kind: str):
    if not meta.introduction.strip():
        return
    _emit_section_title_before(anchor, doc, S.S_PREFACE_TITLE, "引言", kind=kind)
    _emit_introduction_content(anchor, doc, meta)


def _emit_cover_references(anchor: Paragraph, doc, sdoc: model.StandardDoc, kind: str):
    if not sdoc.references:
        return
    _emit_section_title_before(anchor, doc, S.S_REF_TITLE, "参考文献", kind=kind)
    for item in sdoc.references:
        _new_paragraph_before(anchor, doc, S.S_REF_ITEM, text=item)


def _emit_cover_index(anchor: Paragraph, doc, sdoc: model.StandardDoc, kind: str):
    if not sdoc.index_groups:
        return
    _emit_section_title_before(anchor, doc, S.S_INDEX_TITLE, "索引", kind=kind)
    for group in sdoc.index_groups:
        _new_paragraph_before(anchor, doc, S.S_INDEX_LETTER, text=group.letter)
        for item in group.items:
            _emit_index_item(anchor, doc, item)


def _emit_document_end_line(anchor: Paragraph, doc, image_bytes: bytes):
    para = _new_paragraph_before(anchor, doc, S.S_PARA)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(io.BytesIO(image_bytes))


def _emit_cover_sections(doc, sdoc: model.StandardDoc, end_line_image: bytes, kind: str):
    anchor = doc.add_paragraph()
    refs = _body_page_refs(doc)
    include_even = sdoc.meta.odd_even_pages
    body_started = False

    def front_break(start: Optional[int] = None):
        _new_section_break_before(
            anchor,
            doc,
            refs=refs,
            page_fmt="upperRoman",
            page_start=start,
            include_even=include_even,
        )

    def body_break(start: Optional[int] = None):
        nonlocal body_started
        _new_section_break_before(
            anchor,
            doc,
            refs=refs,
            page_start=start,
            include_even=include_even,
        )
        body_started = True

    _emit_cover_toc(anchor, doc, kind)
    front_break(start=1)
    _emit_cover_foreword(anchor, doc, sdoc.meta, kind)
    if sdoc.meta.introduction.strip():
        front_break()
        _emit_cover_introduction(anchor, doc, sdoc.meta, kind)
    front_break()
    _emit_body_standard_title(anchor, doc, sdoc.meta)
    _emit_body_blocks(anchor, doc, sdoc.body)
    for idx, appx in enumerate(sdoc.appendices):
        body_break(start=1 if not body_started else None)
        _emit_appendices(
            anchor,
            doc,
            [appx],
            page_break_before=False,
            section_before_each=False,
            start_index=idx,
        )
    if sdoc.references:
        body_break(start=1 if not body_started else None)
        _emit_cover_references(anchor, doc, sdoc, kind)
    if sdoc.index_groups:
        body_break(start=1 if not body_started else None)
        _emit_cover_index(anchor, doc, sdoc, kind)
    _emit_document_end_line(anchor, doc, end_line_image)
    _configure_final_section(
        doc,
        refs=refs,
        page_start=1 if not body_started else None,
        include_even=include_even,
    )
    _remove_paragraph(anchor)


# --------------------------------------------------------------------------- #
# 入口
# --------------------------------------------------------------------------- #
def build_cover(
    sdoc: model.StandardDoc,
    output_path: str,
    kind: str = "auto",
    cover_form_protection: Optional[bool] = None,
):
    """封面蓝图后端：以封面蓝图为基底，正文全部由代码顺序生成。"""
    _reset_counters()

    resolved_kind = _resolve_kind(kind, sdoc.meta)
    cover_path = _default_cover_path(resolved_kind)

    _copy_cover_base(cover_path, output_path)
    end_line_image = _read_cover_end_line_image(output_path, resolved_kind)

    doc = Document(output_path)
    _configure_standard_styles(doc)

    cover_info = _apply_cover_fields(doc, sdoc.meta, kind=resolved_kind)
    _ensure_cover_publisher(doc, sdoc.meta.publisher, cover_info, kind=resolved_kind)
    _cleanup_cover_placeholders(doc)

    _emit_cover_sections(doc, sdoc, end_line_image, resolved_kind)
    _enable_update_fields(doc)
    _set_even_and_odd_headers(doc, sdoc.meta.odd_even_pages)
    if _should_enable_cover_form_protection(sdoc.meta, cover_form_protection):
        _enable_cover_form_field_protection(doc)
    else:
        _disable_form_field_protection(doc)

    doc.save(output_path)
    return output_path


def build(
    sdoc: model.StandardDoc,
    template_path: str,
    output_path: str,
    kind: str = "auto",
    cover_form_protection: Optional[bool] = None,
):
    _reset_counters()

    # 复制模板再打开，避免改动原模板
    tmp = output_path
    shutil.copyfile(template_path, tmp)
    doc = Document(tmp)

    _configure_standard_styles(doc)

    cover_kind = _resolve_kind(kind, sdoc.meta)
    cover_info = _apply_cover_fields(doc, sdoc.meta, kind=cover_kind)
    _ensure_cover_publisher(doc, sdoc.meta.publisher, cover_info, kind=cover_kind)
    _cleanup_cover_placeholders(doc)
    _build_foreword(doc, sdoc.meta)
    _build_introduction(doc, sdoc.meta)
    _build_body(doc, sdoc)
    _build_appendices(doc, sdoc)
    _build_references(doc, sdoc)
    _build_index(doc, sdoc)
    _build_toc(doc)
    _enable_update_fields(doc)
    _set_even_and_odd_headers(doc, sdoc.meta.odd_even_pages)
    if _should_enable_cover_form_protection(sdoc.meta, cover_form_protection):
        _enable_cover_form_field_protection(doc)
    else:
        _disable_form_field_protection(doc)

    doc.save(output_path)
    return output_path
